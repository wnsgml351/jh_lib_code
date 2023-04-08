import requests
from bs4 import BeautifulSoup
import pyautogui
from docx import Document

# 키워드 입력 받기
keyword = pyautogui.prompt("검색어를 입력하세요.")

# 워드문서 객체 생성하기
document = Document()



response = requests.get(f"https://search.naver.com/search.naver?where=news&sm=tab_jum&query={keyword}")
html = response.text
soup = BeautifulSoup(html, 'html.parser')
articles = soup.select("div.info_group") # 뉴스 기사 div 10개 추출
for article in articles:
    links = article.select("a.info")
    if len(links) >= 2: # 링크가 2개 이상이면
        url = links[1].attrs['href'] # 두번째 링크의 href 추출
        # 다시 request를 날려 준다
        response = requests.get(url, headers={'User-Agent' : 'Mozila/5.0'})
        html = response.text
        soup = BeautifulSoup(html, 'html.parser')
        
        # 연예 뉴스 체크
        if "entertain" in response.url: 
            title = soup.select_one(".end_tit")
            content = soup.select_one("#articeBody")
        elif "sports" in response.url:
            title = soup.select_one("h4.title")
            content = soup.select_one("#newsEndContents")

            # 본문 내용 중에 불 필요한 div 제거
            divs = content.select("div")
            for div in divs:
                div.decompose()
            
            # 본문 내용 중에 불 필요한 p 제거
            paragraphs = content.select("p")
            for p in paragraphs:
                p.decompose()

        else: 
            title = soup.select_one(".media_end_head_headline")
            content = soup.select_one("#newsct_article")

        print("=======링크======= \n", url)
        print("=======제목======= \n", title.text.strip())
        print("=======본문======= \n", content.text.strip())

        # 워드에 제목, 링크, 본문 저장하기
        document.add_heading(title.text.strip(), level=0)
        document.add_paragraph(url)
        document.add_paragraph(content.text.strip())

# 워드문서 저장하기
document.save(f"{keyword}_result.docx")

